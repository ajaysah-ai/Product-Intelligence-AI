from docx import Document as DocxDocument

from app.extraction.errors import ExtractionError
from app.extraction.image_extractor import extract_images_concurrently

MIN_IMAGE_BYTES = 2000  # skip tiny icons/logos — same reasoning as the PDF extractor
MAX_IMAGES_OCR_PER_DOC = 5  # same reasoning as the PDF extractor — bounds worst-case work


def extract_docx(path: str) -> str:
    try:
        d = DocxDocument(path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]

        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        # Embedded pictures live as relationship parts (word/media/*) — python-docx
        # doesn't surface these through the paragraph API, so we go via .part.rels.
        # Collected up front and OCR'd concurrently, same pattern as the PDF extractor.
        ocr_tasks: list[bytes] = []
        for rel in d.part.rels.values():
            if len(ocr_tasks) >= MAX_IMAGES_OCR_PER_DOC:
                break
            if "image" in rel.reltype:
                try:
                    image_bytes = rel.target_part.blob
                    if len(image_bytes) < MIN_IMAGE_BYTES:
                        continue
                    ocr_tasks.append(image_bytes)
                except Exception:
                    continue  # one bad embedded image shouldn't fail the document

        parts.extend(extract_images_concurrently(ocr_tasks))

        return "\n".join(parts)
    except Exception as e:
        raise ExtractionError(f"DOCX extraction failed: {e}") from e
