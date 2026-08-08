"""
Run after the backend is up (this one is slower — EasyOCR downloads its model
on first use, and the large-CSV test writes/reads a real 150k-row file):
    docker compose exec backend python -m app.test_phase3

Exits non-zero if any check fails.
"""
import csv
import io
import sys
import tempfile
import time
from pathlib import Path

import fitz  # PyMuPDF
import psutil
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from PIL import Image, ImageDraw, ImageFont

from app.extraction.dispatcher import extract_file
from app.extraction.errors import ExtractionError
from app.extraction.tabular_extractor import extract_csv
from app.main import app

client = TestClient(app)
results = []


def check(description, passed):
    results.append((description, passed))
    mark = "PASS" if passed else "FAIL"
    print(f"[{mark}] {description}")


def make_pdf(path: Path, text: str):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(str(path))
    doc.close()


def make_docx(path: Path, text: str):
    d = DocxDocument()
    d.add_paragraph(text)
    d.save(str(path))


def make_ocr_image(path: Path, text: str, size=(800, 200), font_size=60):
    img = Image.new("RGB", size, color="white")
    draw = ImageDraw.Draw(img)
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    font = ImageFont.truetype(font_path, font_size)
    draw.text((30, size[1] // 2 - font_size // 2), text, fill="black", font=font)
    img.save(str(path))


def main():
    tmp = Path(tempfile.mkdtemp())

    # -----------------------------------------------------------------
    # 1) One sample file per type -> correct text extracted
    # -----------------------------------------------------------------
    pdf_path = tmp / "sample.pdf"
    make_pdf(pdf_path, "Wireless Mechanical Keyboard Spec Sheet")
    try:
        text = extract_file(str(pdf_path))
        check("PDF: text extracted correctly", "Wireless Mechanical Keyboard" in text)
    except ExtractionError as e:
        check(f"PDF: text extracted correctly (raised {e})", False)

    docx_path = tmp / "sample.docx"
    make_docx(docx_path, "Hot-swappable switches, RGB backlight, USB-C connector")
    try:
        text = extract_file(str(docx_path))
        check("DOCX: text extracted correctly", "Hot-swappable switches" in text)
    except ExtractionError as e:
        check(f"DOCX: text extracted correctly (raised {e})", False)

    txt_path = tmp / "sample.txt"
    txt_path.write_text("Plain text product notes: battery life 40 hours.", encoding="utf-8")
    try:
        text = extract_file(str(txt_path))
        check("TXT: text extracted correctly", "battery life 40 hours" in text)
    except ExtractionError as e:
        check(f"TXT: text extracted correctly (raised {e})", False)

    html_path = tmp / "sample.html"
    html_path.write_text(
        "<html><head><style>body{color:red}</style></head>"
        "<body><h1>Product Page</h1><p>Waterproof rating IP67</p>"
        "<script>console.log('noise')</script></body></html>",
        encoding="utf-8",
    )
    try:
        text = extract_file(str(html_path))
        check(
            "HTML: text extracted, script/style stripped",
            "Waterproof rating IP67" in text and "console.log" not in text,
        )
    except ExtractionError as e:
        check(f"HTML: text extracted correctly (raised {e})", False)

    json_path = tmp / "sample.json"
    json_path.write_text('{"product": "Keyboard", "warranty_years": 2}', encoding="utf-8")
    try:
        text = extract_file(str(json_path))
        check("JSON: text extracted correctly", "warranty_years" in text and "2" in text)
    except ExtractionError as e:
        check(f"JSON: text extracted correctly (raised {e})", False)

    csv_path = tmp / "sample.csv"
    csv_path.write_text("name,price\nKeyboard,4999\nMouse,1299\n", encoding="utf-8")
    try:
        text = extract_file(str(csv_path))
        check("CSV: text extracted correctly", "Keyboard" in text and "4999" in text)
    except ExtractionError as e:
        check(f"CSV: text extracted correctly (raised {e})", False)

    img_path = tmp / "sample.png"
    make_ocr_image(img_path, "WARRANTY")
    try:
        text = extract_file(str(img_path))
        check(
            "Image OCR: readable text recognized",
            "WARRANTY" in text.upper() or "WARRANT" in text.upper(),
        )
    except ExtractionError as e:
        check(f"Image OCR: readable text recognized (raised {e})", False)

    # -----------------------------------------------------------------
    # 2) Corrupted file -> graceful error, no crash
    # -----------------------------------------------------------------
    corrupt_path = tmp / "corrupt.pdf"
    corrupt_path.write_bytes(b"%PDF-1.4\n" + b"\x00\x01\x02garbage-not-a-real-pdf-body" * 20)
    try:
        extract_file(str(corrupt_path))
        check("Corrupted PDF raises ExtractionError (not a crash)", False)
    except ExtractionError:
        check("Corrupted PDF raises ExtractionError (not a crash)", True)
    except Exception as e:
        check(f"Corrupted PDF raises ExtractionError (got unexpected {type(e).__name__})", False)

    # -----------------------------------------------------------------
    # 2b) Embedded images inside PDF/DOCX are also OCR'd, not just skipped
    # -----------------------------------------------------------------
    embedded_img_path = tmp / "embedded.png"
    make_ocr_image(embedded_img_path, "SERIALX9", size=(500, 150), font_size=50)

    pdf_with_image_path = tmp / "sample_with_image.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Product Model XJ-500 Specification")
    page.insert_image(fitz.Rect(72, 120, 400, 220), filename=str(embedded_img_path))
    doc.save(str(pdf_with_image_path))
    doc.close()
    try:
        text = extract_file(str(pdf_with_image_path))
        check("PDF: text layer extracted alongside embedded image", "Model XJ-500" in text)
        check(
            "PDF: embedded image OCR'd (not silently skipped)",
            "SERIALX9" in text.upper() or "SERIAL" in text.upper(),
        )
    except ExtractionError as e:
        check(f"PDF with embedded image (raised {e})", False)

    docx_with_image_path = tmp / "sample_with_image.docx"
    d = DocxDocument()
    d.add_paragraph("Product Manual Cover Page")
    d.add_picture(str(embedded_img_path))
    d.save(str(docx_with_image_path))
    try:
        text = extract_file(str(docx_with_image_path))
        check("DOCX: paragraph text extracted alongside embedded image", "Product Manual Cover Page" in text)
        check(
            "DOCX: embedded image OCR'd (not silently skipped)",
            "SERIALX9" in text.upper() or "SERIAL" in text.upper(),
        )
    except ExtractionError as e:
        check(f"DOCX with embedded image (raised {e})", False)

    # -----------------------------------------------------------------
    # 3) 5 files together via the real API -> all extracted, one failure isolated
    # -----------------------------------------------------------------
    r = client.post(
        "/submit",
        data={"text": "batch extraction test"},
        files=[
            ("files", ("batch1.txt", io.BytesIO(b"batch file one content"), "text/plain")),
            ("files", ("batch2.txt", io.BytesIO(b"batch file two content"), "text/plain")),
            ("files", ("batch3.csv", io.BytesIO(b"a,b\n1,2\n"), "text/csv")),
            ("files", ("batch4.json", io.BytesIO(b'{"k": "v"}'), "application/json")),
            (
                "files",
                (
                    "batch5.html",
                    io.BytesIO(b"<html><body>hello</body></html>"),
                    "text/html",
                ),
            ),
        ],
    )
    check("batch submit (5 files) -> 200", r.status_code == 200)
    request_id = r.json().get("request_id") if r.status_code == 200 else None

    if request_id:
        r2 = client.post(f"/extract/{request_id}")
        check("batch extract -> 200", r2.status_code == 200)
        if r2.status_code == 200:
            statuses = [item["status"] for item in r2.json()["results"]]
            check("batch extract -> all 5 files succeeded", statuses == ["success"] * 5)
    else:
        check("batch extract -> 200 (skipped, no request_id)", False)

    # -----------------------------------------------------------------
    # 4) Large CSV (150k rows) -> completes fast, memory stays bounded
    # -----------------------------------------------------------------
    large_csv_path = tmp / "large.csv"
    num_rows = 150_000
    with open(large_csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["sku", "name", "price"])
        for i in range(num_rows):
            writer.writerow([f"SKU{i}", f"Product {i}", i * 1.5])

    file_size_mb = large_csv_path.stat().st_size / (1024 * 1024)
    process = psutil.Process()
    mem_before = process.memory_info().rss

    start = time.time()
    text = extract_csv(str(large_csv_path))
    elapsed = time.time() - start

    mem_after = process.memory_info().rss
    mem_growth_mb = (mem_after - mem_before) / (1024 * 1024)

    line_count = text.count("\n") + 1
    check(
        f"large CSV ({num_rows} rows, {file_size_mb:.1f}MB): all rows present",
        line_count == num_rows + 1,  # +1 header
    )
    check(f"large CSV: completed in {elapsed:.2f}s (< 15s)", elapsed < 15)
    check(
        f"large CSV: memory growth {mem_growth_mb:.1f}MB stayed under 5x file size ({file_size_mb * 5:.1f}MB)",
        mem_growth_mb < file_size_mb * 5,
    )

    total = len(results)
    passed = sum(1 for _, ok in results if ok)
    print(f"\n{passed}/{total} checks passed")
    if passed != total:
        sys.exit(1)


if __name__ == "__main__":
    main()
